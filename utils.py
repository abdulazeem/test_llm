import os
from typing import List
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.document_loaders import Docx2txtLoader
from langchain.document_loaders import TextLoader
from langchain.document_loaders.csv_loader import CSVLoader
import yaml
import io
import markdown
from docx import Document
from docx.shared import Pt
from lxml import html
from googletrans import Translator
from langdetect import detect


def markdown_to_docx(markdown_text):
    markdown_text = markdown_text.strip()
    pre_text = '\n'.join([text.strip() for text in markdown_text.split('\n')])

    html_content = markdown.markdown(pre_text)

    # Create a new Document
    doc = Document()

    # Parse the HTML content
    md_elements = html.fragments_fromstring(html_content)

    # Iterate through the Markdown elements and insert into the DOCX document
    for md_element in md_elements:
        if md_element.tag == 'p':
            doc.add_paragraph(md_element.text)
        elif md_element.tag == 'h1':
            doc.add_heading(md_element.text, level=1)
        elif md_element.tag == 'h2':
            doc.add_heading(md_element.text, level=2)
        elif md_element.tag == 'h3':
            doc.add_heading(md_element.text, level=3)
        elif md_element.tag == 'em':
            run = doc.add_paragraph().add_run(md_element.text)
            run.italic = True
        elif md_element.tag == 'strong':
            run = doc.add_paragraph().add_run(md_element.text)
            run.bold = True
        elif md_element.tag == 'ul':
            for li in md_element.iterdescendants('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif md_element.tag == 'ol':
            for li in md_element.iterdescendants('li'):
                doc.add_paragraph(li.text, style='List Number')
        elif md_element.tag == 'pre':
            code_block = doc.add_paragraph().add_run(md_element.text)
            code_block.font.size = Pt(10)
            code_block.font.name = 'Courier New'
    # Save the document to the specified output path
    return doc

def detect_and_translate(query):
    detected_language = detect(query)
    translator = Translator()
    translated_text = translator.translate(query, dest='en')
    return detected_language, translated_text.text

def translate_response(response, target_language):
     translator = Translator()
     translated_response = translator.translate(response, dest=target_language)
     return translated_response.text

def load_yaml(name):
     with open(name) as fp:
        config = yaml.safe_load(fp)
     return config

def load_doc_from_dir(dirpath:str):
     documents = []
     for file in os.listdir(dirpath):
        if file.lower().endswith('.pdf'):
            pdf_path = os.path.join(dirpath, file)
            loader = PyPDFLoader(pdf_path)
            documents.extend(loader.load())
        elif file.lower().endswith('.docx') or file.endswith('.doc'):
            doc_path = os.path.join(dirpath, file)
            loader = Docx2txtLoader(doc_path)
            documents.extend(loader.load())
        elif file.lower().endswith('.txt'):
            text_path = os.path.join(dirpath, file)
            loader = TextLoader(text_path)
            documents.extend(loader.load())
            documents.extend(loader.load())
     # text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100)
     # text_splitter =RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100, separators=["\n\n", "\n", " ", ""])
     # text_splitter = CharacterTextSplitter(separator='\n', chunk_size=1000, chunk_overlap=100)
     text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
     chunked_documents = text_splitter.split_documents(documents)
     print('**************************************************')
     # print(len(chunked_documents))
     print(f"Splitted into {len(chunked_documents)}")
     print('**************************************************')
     print('**************************************************')
     return chunked_documents